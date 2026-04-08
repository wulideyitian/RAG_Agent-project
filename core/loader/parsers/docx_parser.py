"""
Word 文件解析器（增强版 - 支持标题层级提取）
"""
from typing import List, Dict, Optional
from langchain_core.documents import Document
from langchain_community.document_loaders.word_document import Docx2txtLoader
from app.utils.logger_handler import logger
from core.preprocessing.text_cleaner import TextCleaner
import re

try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    logger.warning("未安装 python-docx，将使用基础解析模式")


class DOCXParser:
    """Word 文件解析器（增强版 - 支持标题层级提取）"""
    
    def __init__(self, enable_structure_extraction: bool = True):
        """
        :param enable_structure_extraction: 是否启用结构提取（标题层级）
        """
        self.text_cleaner = TextCleaner()
        self.enable_structure_extraction = enable_structure_extraction and HAS_PYTHON_DOCX
        if self.enable_structure_extraction:
            logger.info("Word 结构提取已启用")
    
    def parse(self, filepath: str) -> List[Document]:
        """
        解析 Word 文件
        :param filepath: 文件路径
        :return: Document 列表
        """
        try:
            # 如果启用了结构提取，使用增强解析
            if self.enable_structure_extraction:
                return self._parse_with_structure(filepath)
            else:
                # 否则使用基础解析
                return self._parse_basic(filepath)
            
        except Exception as e:
            logger.error(f"解析 Word 文件{filepath}失败：{str(e)}")
            raise Exception(f"Word 解析失败：{str(e)}")
    
    def _parse_basic(self, filepath: str) -> List[Document]:
        """
        基础解析：仅提取纯文本
        :param filepath: 文件路径
        :return: Document 列表
        """
        loader = Docx2txtLoader(filepath)
        documents = loader.load()
        
        if not documents:
            return []
        
        # 清理文本
        cleaned_content = self.text_cleaner.clean(
            documents[0].page_content, 
            'docx'
        )
        
        doc = Document(
            page_content=cleaned_content, 
            metadata={"source": filepath}
        )
        
        logger.info(f"Word 文件{filepath}基础解析成功")
        return [doc]
    
    def _parse_with_structure(self, filepath: str) -> List[Document]:
        """
        增强解析：提取标题层级结构
        :param filepath: 文件路径
        :return: Document 列表（按标题分块）
        """
        try:
            doc = DocxDocument(filepath)
        except Exception as e:
            logger.warning(f"无法使用 python-docx 解析，回退到基础模式: {e}")
            return self._parse_basic(filepath)
        
        # 提取结构化内容
        structured_chunks = self._extract_structured_chunks(doc)
        
        if not structured_chunks:
            logger.warning(f"Word 文件{filepath}未提取到结构化内容")
            return []
        
        # 转换为 Document 列表
        documents = []
        for chunk in structured_chunks:
            # 清理文本
            cleaned_content = self.text_cleaner.clean(chunk['content'], 'docx')
            
            if not cleaned_content or not cleaned_content.strip():
                continue
            
            doc = Document(
                page_content=cleaned_content,
                metadata={
                    "source": filepath,
                    "heading_level": chunk.get('heading_level', 0),
                    "heading_text": chunk.get('heading_text', ''),
                    "section_index": chunk.get('section_index', 0),
                }
            )
            documents.append(doc)
        
        logger.info(f"Word 文件{filepath}增强解析成功，共{len(documents)}个结构块")
        return documents
    
    def _extract_structured_chunks(self, doc: DocxDocument) -> List[Dict]:
        """
        从 Word 文档中提取结构化块（按标题分割）
        :param doc: python-docx Document 对象
        :return: 结构化块列表
        """
        chunks = []
        current_heading_level = 0
        current_heading_text = ""
        current_content = []
        section_index = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # 检测是否为标题
            heading_level = self._detect_heading_level(paragraph)
            
            if heading_level > 0:
                # 保存之前的块
                if current_content:
                    chunks.append({
                        'heading_level': current_heading_level,
                        'heading_text': current_heading_text,
                        'content': '\n'.join(current_content),
                        'section_index': section_index
                    })
                    section_index += 1
                
                # 开始新块
                current_heading_level = heading_level
                current_heading_text = text
                current_content = []
            else:
                # 添加到当前块的内容
                current_content.append(text)
        
        # 保存最后一个块
        if current_content:
            chunks.append({
                'heading_level': current_heading_level,
                'heading_text': current_heading_text,
                'content': '\n'.join(current_content),
                'section_index': section_index
            })
        
        logger.debug(f"提取到 {len(chunks)} 个结构化块")
        return chunks
    
    def _detect_heading_level(self, paragraph) -> int:
        """
        检测段落的标题层级
        :param paragraph: python-docx Paragraph 对象
        :return: 标题层级 (0 表示非标题, 1-9 表示 Heading 1-9)
        """
        # 方法1: 检查样式名称
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        if 'heading' in style_name:
            # 提取数字，如 "Heading 1" -> 1
            match = re.search(r'heading\s*(\d+)', style_name)
            if match:
                return int(match.group(1))
        
        # 方法2: 检查字体大小和加粗（启发式方法）
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_size = first_run.font.size
            is_bold = first_run.bold
            
            # 大号字体且加粗可能是标题
            if is_bold and font_size and font_size.pt >= 14:
                return 1  # 假设为主标题
            elif is_bold and font_size and font_size.pt >= 12:
                return 2  # 假设为副标题
        
        return 0
